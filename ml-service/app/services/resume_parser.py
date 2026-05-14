import io
import os
from dataclasses import dataclass

import pdfplumber
from docx import Document

from app.utils.preprocessing import normalize_text


class ResumeParseError(Exception):
    pass


@dataclass(frozen=True)
class ParsedResumeText:
    text: str
    plain_text: str
    page_count: int | None = None
    word_count: int = 0


def _extract_pdf_text(file_bytes: bytes) -> tuple[str, int]:
    text_chunks: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text(x_tolerance=1, y_tolerance=3) or ""
            if page_text.strip():
                text_chunks.append(page_text)
        return "\n\n".join(text_chunks), len(pdf.pages)


def _extract_docx_text(file_bytes: bytes) -> str:
    document = Document(io.BytesIO(file_bytes))
    lines: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))

    return "\n".join(lines)


def parse_resume_text(file_bytes: bytes, filename: str) -> ParsedResumeText:
    extension = os.path.splitext((filename or "").lower())[1]
    page_count: int | None = None

    if extension == ".pdf":
        raw_text, page_count = _extract_pdf_text(file_bytes)
    elif extension == ".docx":
        raw_text = _extract_docx_text(file_bytes)
    else:
        raise ResumeParseError("Unsupported file format. Only PDF and DOCX are allowed.")

    cleaned = normalize_text(raw_text, preserve_lines=True)
    plain_text = normalize_text(raw_text)
    if not plain_text:
        raise ResumeParseError(
            "No readable text found in resume. If this is a scanned PDF, export it with selectable text before uploading."
        )

    return ParsedResumeText(
        text=cleaned,
        plain_text=plain_text,
        page_count=page_count,
        word_count=len(plain_text.split()),
    )
